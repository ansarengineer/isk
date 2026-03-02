# -*- coding: utf-8 -*-
"""
Справки-бот: берет данные из Excel, подставляет в Word-шаблон, формирует справки
по одному ответчику на строку.

Требования:
- «Дата_составления_иска» = дата создания справки (сегодняшняя дата)
- Весь текст итогового DOCX = Times New Roman, 12 pt
- 4 шаблона, выбор по столбцам:
    - "Истец" (45)
    - "Путь к шаблону по Справкам" (126)
- Название файла справки берём из столбца "Справка" (123). Если нет/пусто — fallback.
- GUI:
    1) загрузка Excel
    2) загрузка 4 шаблонов
    3) выбор папки сохранения
- Авто-папка результатов: подпапка с названием Excel-файла (без расширения)

Зависимости:
    pip install pandas openpyxl python-docx
"""

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Tuple

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.oxml.ns import qn

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except Exception:
    tk = None


PLACEHOLDER_RE = re.compile(r"«([^»]+)»")  # «КЛЮЧ»


# ---------------- utils ----------------

def normalize_key(col_name: str) -> str:
    """Excel-колонка -> ключ плейсхолдера: пробелы -> _"""
    s = str(col_name).strip()
    s = re.sub(r"\s+", "_", s)
    return s


def safe_str(v) -> str:
    if pd.isna(v):
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v).strip()


def sanitize_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r'[\\/:*?"<>|]+', "_", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_placeholders_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str]) -> None:
    txt = paragraph.text
    if "«" not in txt:
        return

    def repl(m):
        key = m.group(1)
        return mapping.get(key, m.group(0))

    new_txt = PLACEHOLDER_RE.sub(repl, txt)
    if new_txt != txt:
        paragraph.text = new_txt


def replace_all(doc: Document, mapping: Dict[str, str]) -> None:
    for p in iter_paragraphs(doc):
        replace_placeholders_in_paragraph(p, mapping)


def enforce_times_new_roman_12(doc: Document) -> None:
    """Принудительно ставит Times New Roman 12pt для всего текста (включая таблицы)."""
    def apply(paragraph: Paragraph):
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for p in doc.paragraphs:
        apply(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply(p)


# ---------------- selection logic ----------------

def detect_template_kind(istec: str, path_hint: str) -> str:
    """
    Возвращает: 'ep', 'usmanova', 'kaipov', 'zhag'
    Приоритет: "Путь к шаблону по Справкам" -> "Истец".
    """
    istec_l = (istec or "").lower()
    hint_l = (path_hint or "").lower()

    if hint_l:
        if "эп" in hint_l or "эксперт" in hint_l or "expert" in hint_l or "ep" in hint_l:
            return "ep"
        if "усман" in hint_l:
            return "usmanova"
        if "каип" in hint_l:
            return "kaipov"
        if "жаг" in hint_l or "жагып" in hint_l:
            return "zhag"

    if "эксперт плюс" in istec_l or "коллекторское агентство" in istec_l:
        return "ep"
    if "усманова асель" in istec_l:
        return "usmanova"
    if "каипов ришат" in istec_l:
        return "kaipov"
    if "жагыпарова" in istec_l or "гүлжан" in istec_l:
        return "zhag"

    return "ep"


def build_mapping_from_row(row: pd.Series) -> Dict[str, str]:
    """
    mapping:
      normalize_key(col) -> value
    чтобы "Дата составления иска" => "Дата_составления_иска"
    """
    mapping: Dict[str, str] = {}
    for col in row.index:
        mapping[normalize_key(col)] = safe_str(row[col])
    return mapping


def make_spravka_filename(row: pd.Series, idx: int) -> str:
    """
    Имя справки:
    - Если заполнен столбец "Справка" (123) -> используем его (как есть, +.docx если нет)
    - Иначе fallback: Справка_<ФИО>_<номер_договора>.docx
    """
    desired = safe_str(row.get("Справка", ""))
    desired = sanitize_filename(desired)

    if desired:
        if desired.lower().endswith(".docx"):
            return desired
        return desired + ".docx"

    # fallback
    fio = safe_str(row.get("ФИО", "")) or f"Ответчик_{idx+1}"
    dog = safe_str(row.get("номер договора", "")) or safe_str(row.get("номер_договора", ""))

    fio = sanitize_filename(fio)[:80]
    dog = sanitize_filename(dog)[:30]

    parts = ["Справка", fio]
    if dog:
        parts.append(dog)
    return "_".join([p for p in parts if p]) + ".docx"


# ---------------- core ----------------

@dataclass
class Inputs:
    excel_path: Path
    template_ep: Path
    template_usmanova: Path
    template_kaipov: Path
    template_zhag: Path
    out_root: Path  # корневая папка, внутри создадим подпапку по имени excel


def generate_spravki(inputs: Inputs) -> Tuple[int, Path]:
    df = pd.read_excel(inputs.excel_path)
    if df.empty:
        raise ValueError("Excel пустой — нечего формировать.")

    out_dir = inputs.out_root / (sanitize_filename(inputs.excel_path.stem) or "Справки")
    out_dir.mkdir(parents=True, exist_ok=True)

    templates = {
        "ep": inputs.template_ep,
        "usmanova": inputs.template_usmanova,
        "kaipov": inputs.template_kaipov,
        "zhag": inputs.template_zhag,
    }

    today_str = datetime.now().strftime("%d.%m.%Y")
    created = 0

    for i, (_, row) in enumerate(df.iterrows()):
        istec = safe_str(row.get("Истец", ""))
        path_hint = safe_str(row.get("Путь к шаблону по Справкам", ""))

        kind = detect_template_kind(istec, path_hint)
        template_path = templates.get(kind, inputs.template_ep)

        doc = Document(template_path)
        mapping = build_mapping_from_row(row)

        # дата создания справки
        mapping["Дата_составления_иска"] = today_str

        replace_all(doc, mapping)
        enforce_times_new_roman_12(doc)

        out_name = make_spravka_filename(row, i)
        doc.save(out_dir / out_name)

        created += 1

    return created, out_dir


# ---------------- GUI ----------------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Справки-бот (Excel → Word)")
        self.resizable(False, False)

        self.excel_path: Optional[Path] = None
        self.out_root: Optional[Path] = None

        self.template_ep: Optional[Path] = None
        self.template_usmanova: Optional[Path] = None
        self.template_kaipov: Optional[Path] = None
        self.template_zhag: Optional[Path] = None

        pad = {"padx": 10, "pady": 6}

        tk.Label(self, text="1) Excel с данными:").grid(row=0, column=0, sticky="w", **pad)
        tk.Button(self, text="Загрузить Excel…", command=self.pick_excel, width=24).grid(row=0, column=1, **pad)
        self.val_excel = tk.Label(self, text="—", fg="#555", wraplength=560, justify="left")
        self.val_excel.grid(row=1, column=0, columnspan=2, sticky="w", **pad)

        tk.Label(self, text="2) 4 шаблона справок (.docx):").grid(row=2, column=0, sticky="w", **pad)

        tk.Button(self, text="Шаблон ЭП…", command=self.pick_ep, width=24).grid(row=3, column=0, **pad)
        self.val_ep = tk.Label(self, text="—", fg="#555", wraplength=560, justify="left")
        self.val_ep.grid(row=3, column=1, sticky="w", **pad)

        tk.Button(self, text="Шаблон Усманова…", command=self.pick_usmanova, width=24).grid(row=4, column=0, **pad)
        self.val_usm = tk.Label(self, text="—", fg="#555", wraplength=560, justify="left")
        self.val_usm.grid(row=4, column=1, sticky="w", **pad)

        tk.Button(self, text="Шаблон Каипов…", command=self.pick_kaipov, width=24).grid(row=5, column=0, **pad)
        self.val_kai = tk.Label(self, text="—", fg="#555", wraplength=560, justify="left")
        self.val_kai.grid(row=5, column=1, sticky="w", **pad)

        tk.Button(self, text="Шаблон Жаг…", command=self.pick_zhag, width=24).grid(row=6, column=0, **pad)
        self.val_zh = tk.Label(self, text="—", fg="#555", wraplength=560, justify="left")
        self.val_zh.grid(row=6, column=1, sticky="w", **pad)

        tk.Label(self, text="3) Куда сохранить (корневая папка):").grid(row=7, column=0, sticky="w", **pad)
        tk.Button(self, text="Выбрать папку…", command=self.pick_out_root, width=24).grid(row=7, column=1, **pad)
        self.val_out = tk.Label(self, text="—", fg="#555", wraplength=560, justify="left")
        self.val_out.grid(row=8, column=0, columnspan=2, sticky="w", **pad)

        tk.Button(self, text="Сформировать справки", command=self.run, width=34).grid(
            row=9, column=0, columnspan=2, pady=12
        )

        self.attributes("-topmost", True)
        self.after(350, lambda: self.attributes("-topmost", False))

    def pick_excel(self):
        path = filedialog.askopenfilename(
            title="Выберите Excel",
            filetypes=[("Excel files", "*.xlsx *.xls *.xlsm *.xlsb"), ("All files", "*.*")]
        )
        if path:
            self.excel_path = Path(path)
            self.val_excel.config(text=str(self.excel_path))

    def _pick_docx(self, title: str) -> Optional[Path]:
        path = filedialog.askopenfilename(
            title=title,
            filetypes=[("Word files", "*.docx *.docm"), ("All files", "*.*")]
        )
        return Path(path) if path else None

    def pick_ep(self):
        p = self._pick_docx("Выберите шаблон справки ЭП")
        if p:
            self.template_ep = p
            self.val_ep.config(text=str(p))

    def pick_usmanova(self):
        p = self._pick_docx("Выберите шаблон справки Усманова")
        if p:
            self.template_usmanova = p
            self.val_usm.config(text=str(p))

    def pick_kaipov(self):
        p = self._pick_docx("Выберите шаблон справки Каипов")
        if p:
            self.template_kaipov = p
            self.val_kai.config(text=str(p))

    def pick_zhag(self):
        p = self._pick_docx("Выберите шаблон справки Жаг")
        if p:
            self.template_zhag = p
            self.val_zh.config(text=str(p))

    def pick_out_root(self):
        path = filedialog.askdirectory(title="Выберите корневую папку для сохранения")
        if path:
            self.out_root = Path(path)
            self.val_out.config(text=str(self.out_root))

    def run(self):
        if not self.excel_path or not self.out_root:
            messagebox.showwarning("Не хватает данных", "Выберите Excel и папку сохранения.")
            return

        if not all([self.template_ep, self.template_usmanova, self.template_kaipov, self.template_zhag]):
            messagebox.showwarning("Не хватает шаблонов", "Выберите все 4 шаблона справок.")
            return

        try:
            total, out_dir = generate_spravki(
                Inputs(
                    excel_path=self.excel_path,
                    template_ep=self.template_ep,
                    template_usmanova=self.template_usmanova,
                    template_kaipov=self.template_kaipov,
                    template_zhag=self.template_zhag,
                    out_root=self.out_root,
                )
            )
            messagebox.showinfo("Готово", f"Сформировано справок: {total}\nПапка: {out_dir}")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))


def main():
    if tk is None:
        raise RuntimeError("Tkinter не доступен. Запустите на Windows/macOS/Linux с GUI.")
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()
