# -*- coding: utf-8 -*-
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation

import pandas as pd
from docx import Document

# ---------- утилиты ----------

PLACEHOLDER_RE = re.compile(r'«([^»]+)»')

def norm_key(s: str) -> str:
    s = str(s).strip()
    s = s.replace('\n', ' ')
    s = re.sub(r'[«»"]', '', s)
    s = re.sub(r'[\s\-]+', '_', s)
    s = re.sub(r'[^\w_]+', '', s, flags=re.UNICODE)
    return s.lower()

def safe_filename(name: str, max_len: int = 120) -> str:
    name = re.sub(r'[\\/:*?"<>|]+', '_', str(name))
    name = re.sub(r'\s+', ' ', name).strip()
    return name[:max_len] or "иск"

def parse_decimal(value) -> Decimal | None:
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    s = s.replace(' ', '').replace(',', '.')
    try:
        return Decimal(s)
    except InvalidOperation:
        return None

def format_int_money(val: Decimal) -> str:
    v = val.quantize(Decimal('1'), rounding=ROUND_HALF_UP)
    return f"{int(v):,}".replace(",", " ")

def read_table(path: str) -> pd.DataFrame:
    ext = Path(path).suffix.lower()
    if ext in ('.csv', '.txt'):
        try:
            return pd.read_csv(path, sep=';', dtype=str, keep_default_na=False, encoding="utf-8-sig")
        except Exception:
            return pd.read_csv(path, sep=',', dtype=str, keep_default_na=False, encoding="utf-8-sig")
    return pd.read_excel(path, dtype=str, keep_default_na=False)

def collect_placeholders(doc: Document) -> set[str]:
    ph = set()
    for p in doc.paragraphs:
        ph.update(PLACEHOLDER_RE.findall(p.text))
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    ph.update(PLACEHOLDER_RE.findall(p.text))
    return ph

def replace_in_paragraph(paragraph, repl: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    full = ''.join(r.text for r in paragraph.runs)
    new = full
    for k, v in repl.items():
        new = new.replace(f'«{k}»', v)
    if new != full:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ''

def replace_in_doc(doc: Document, repl: dict[str, str]) -> None:
    for p in doc.paragraphs:
        replace_in_paragraph(p, repl)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    replace_in_paragraph(p, repl)

def uniquify(path: Path) -> Path:
    if not path.exists():
        return path
    i = 2
    while True:
        p = path.with_stem(f"{path.stem} ({i})")
        if not p.exists():
            return p
        i += 1

def split_isk_path(val: str):
    parts = re.split(r"[\\/]+", val.strip())
    parts = [p for p in parts if p.strip()]
    if not parts:
        return [], ""
    return [safe_filename(p) for p in parts[:-1]], safe_filename(parts[-1])

# ---------- бизнес-логика заполнения ----------

DEF_FIELDS = ("ФИО", "ИИН", "Адрес")  # что считаем "ответчиком"

def get_cell(row: dict, col_map: dict[str, str], placeholder_name: str) -> str:
    """
    Возвращает значение по плейсхолдеру (ищем колонку с таким же именем).
    Числа округляем до целых.
    """
    col = col_map.get(norm_key(placeholder_name))
    raw = row.get(col, '') if col else ''
    dec = parse_decimal(raw)
    if dec is not None:
        return format_int_money(dec)
    return str(raw) if raw is not None else ''

def get_cell_from_group(rows: list[dict], col_map: dict[str, str], placeholder_name: str) -> str:
    """
    Берёт значение по плейсхолдеру НЕ только из первой строки группы,
    а первое непустое по всем строкам данного "Иска".
    Это важно, когда суммы/истец/госпошлина заполнены не в первой строке.
    """
    col = col_map.get(norm_key(placeholder_name))
    if not col:
        return ""
    for r in rows:
        raw = r.get(col, "")
        if raw is None:
            continue
        s = str(raw).strip()
        if s != "":
            dec = parse_decimal(s)
            return format_int_money(dec) if dec is not None else s
    return ""

def build_replacements_for_claim(rows: list[dict], placeholders: set[str], max_defendants: int = 10) -> dict[str, str]:
    """
    rows: список строк (dict) одного "Иска" в порядке следования.
    1-я строка -> «ФИО», «ИИН», «Адрес»
    2-я строка -> «ФИО1», «ИИН1», «Адрес1»
    ...
    """
    if not rows:
        return {}

    # Карта колонок по первой строке (обычно структура одинакова)
    col_map = {norm_key(k): k for k in rows[0].keys()}

    out: dict[str, str] = {}

    # 1) Общие плейсхолдеры: берём первое непустое значение по группе
    for ph in placeholders:
        nph = norm_key(ph)

        # пропускаем ответчиков — ниже заполним строго по правилам
        if nph in (norm_key(x) for x in DEF_FIELDS):
            continue
        if re.fullmatch(r"(фио|иин|ин|адрес)\d+", nph):
            continue
        if nph in ("фио", "иин", "ин", "адрес"):
            continue

        out[ph] = get_cell_from_group(rows, col_map, ph)

    # 2) Ответчики: максимум 10 строк
    take = rows[:max_defendants]

    for idx, r in enumerate(take):
        suffix = "" if idx == 0 else str(idx)

        for base in DEF_FIELDS:
            ph_name = f"{base}{suffix}"  # «ФИО» или «ФИО1»...
            if ph_name in placeholders:
                out[ph_name] = get_cell(r, col_map, base)

    return out

# ---------- GUI ----------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Генератор исков")
        self.geometry("860x480")

        self.excel = tk.StringVar()
        self.template = tk.StringVar()
        self.output = tk.StringVar()

        self._ui()

    def _ui(self):
        pad = dict(padx=12, pady=8)

        f1 = ttk.LabelFrame(self, text="1) Таблица")
        f1.pack(fill="x", **pad)
        ttk.Entry(f1, textvariable=self.excel).pack(side="left", fill="x", expand=True, padx=10)
        ttk.Button(f1, text="Excel / CSV", command=self.pick_excel).pack(side="left", padx=10)

        f2 = ttk.LabelFrame(self, text="2) Шаблон DOCX")
        f2.pack(fill="x", **pad)
        ttk.Entry(f2, textvariable=self.template).pack(side="left", fill="x", expand=True, padx=10)
        ttk.Button(f2, text="DOCX", command=self.pick_template).pack(side="left", padx=10)

        f3 = ttk.LabelFrame(self, text="3) Папка вывода")
        f3.pack(fill="x", **pad)
        ttk.Entry(f3, textvariable=self.output).pack(side="left", fill="x", expand=True, padx=10)
        ttk.Button(f3, text="Выбрать", command=self.pick_output).pack(side="left", padx=10)

        f4 = ttk.Frame(self)
        f4.pack(fill="x", **pad)
        self.run_btn = ttk.Button(f4, text="Сгенерировать", command=self.run)
        self.run_btn.pack(side="left")
        self.progress = ttk.Progressbar(f4, mode="determinate")
        self.progress.pack(side="left", fill="x", expand=True, padx=10)

        self.log = tk.Text(self, height=10)
        self.log.pack(fill="both", expand=True, padx=12, pady=8)

    def log_msg(self, msg):
        self.log.insert("end", msg + "\n")
        self.log.see("end")
        self.update_idletasks()

    def pick_excel(self):
        p = filedialog.askopenfilename(filetypes=[
            ("Excel/CSV", "*.xlsx *.xls *.csv *.txt"),
            ("All files", "*.*")
        ])
        if p:
            self.excel.set(p)
            self.log_msg(f"Таблица: {p}")

    def pick_template(self):
        p = filedialog.askopenfilename(filetypes=[("DOCX", "*.docx")])
        if p:
            self.template.set(p)
            self.log_msg(f"Шаблон: {p}")

    def pick_output(self):
        p = filedialog.askdirectory()
        if p:
            self.output.set(p)
            self.log_msg(f"Вывод: {p}")

    def run(self):
        if not self.excel.get() or not self.template.get():
            messagebox.showerror("Ошибка", "Выберите таблицу и шаблон")
            return

        try:
            df = read_table(self.excel.get())
        except Exception as e:
            messagebox.showerror("Ошибка чтения", f"Не удалось прочитать таблицу:\n{e}")
            return

        # колонка "Иск" или 122-я
        if "Иск" in df.columns:
            isk_col = "Иск"
        else:
            if len(df.columns) < 122:
                messagebox.showerror(
                    "Ошибка",
                    "Нет колонки 'Иск' и таблица содержит меньше 122 столбцов.\n"
                    "Добавьте колонку 'Иск' или проверьте структуру файла."
                )
                return
            isk_col = df.columns[121]

        out_root = Path(self.output.get() or Path(self.excel.get()).parent)
        out_dir = out_root / "иски"
        out_dir.mkdir(exist_ok=True)

        try:
            base_doc = Document(self.template.get())
            placeholders = collect_placeholders(base_doc)
        except Exception as e:
            messagebox.showerror("Ошибка шаблона", f"Не удалось открыть DOCX:\n{e}")
            return

        groups = list(df.groupby(isk_col, sort=False))

        self.progress["maximum"] = len(groups)
        self.progress["value"] = 0

        self.log_msg(f"Колонка 'Иск': {isk_col}")
        self.log_msg(f"Групп (документов): {len(groups)}")
        self.log_msg(f"Плейсхолдеров в шаблоне: {len(placeholders)}")
        self.log_msg("Правило: 1-я строка = «ФИО», 2-я = «ФИО1», ... максимум 10 ответчиков")
        self.log_msg("Общие поля: берём первое непустое значение по всем строкам одного 'Иска'")

        ok = 0

        for i, (isk_val, gdf) in enumerate(groups, 1):
            try:
                rows = [r.to_dict() for _, r in gdf.iterrows()]

                repl = build_replacements_for_claim(rows, placeholders, max_defendants=10)

                doc = Document(self.template.get())
                replace_in_doc(doc, repl)

                isk_name = str(isk_val).strip()
                if not isk_name:
                    isk_name = f"Иск_{i}"

                subs, name = split_isk_path(isk_name)
                if not name.lower().endswith(".docx"):
                    name += ".docx"

                target = out_dir
                for s in subs:
                    target /= s
                target.mkdir(parents=True, exist_ok=True)

                path = uniquify(target / name)
                doc.save(path)

                ok += 1
                self.progress["value"] = i

                if len(rows) > 10:
                    self.log_msg(f"[{isk_name}] Ответчиков {len(rows)} — взято только 10 (ФИО..ФИО9)")

                if i % 10 == 0:
                    self.log_msg(f"Готово: {i}/{len(groups)}")

            except Exception as e:
                self.log_msg(f"[Группа {i}] Ошибка: {e}")
                self.progress["value"] = i
                continue

        messagebox.showinfo("Готово", f"Сгенерировано документов: {ok} из {len(groups)}")

if __name__ == "__main__":
    App().mainloop()
