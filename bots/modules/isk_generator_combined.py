# -*- coding: utf-8 -*-
"""
Иски-бот: берет данные из Excel, подставляет в Word-шаблон, формирует пачку исков.

Доработки:
- Создаёт подпапку для результатов с названием Excel-файла (без расширения).
  Например: Данные.xlsx -> .../Данные/ (туда сохраняются иски + Реестр_исков.xlsx)

- Исправление: общий текст "По всем вышеуказанным договорам..." НЕ удаляется
  при сокращении количества ответчиков (<10). Границу удаления секций договоров
  ограничиваем ДО этого общего блока.

Зависимости:
    pip install pandas openpyxl python-docx
"""

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import pandas as pd
from docx import Document
from docx.table import _Row
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.oxml.ns import qn

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except Exception:
    tk = None  # если кто-то запускает без GUI


PLACEHOLDER_RE = re.compile(r"«([^»]+)»")  # «КЛЮЧ»


# ---------------- ШРИФТ ----------------

def enforce_times_new_roman_12(doc: Document) -> None:
    """
    Принудительно устанавливает Times New Roman, 12 pt
    для всех абзацев и таблиц документа.
    """
    def apply(paragraph: Paragraph):
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

            # фикс для Word: прописываем все rFonts
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            for k in ("ascii", "hAnsi", "cs", "eastAsia"):
                rFonts.set(qn(f"w:{k}"), "Times New Roman")

    for p in doc.paragraphs:
        apply(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply(p)


# ---------------- ОБЩИЕ УТИЛИТЫ ----------------

def normalize_key(col_name: str) -> str:
    s = str(col_name).strip()
    s = re.sub(r"\s+", "_", s)
    return s


def safe_str(v) -> str:
    if pd.isna(v):
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v).strip()


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_placeholders_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str]) -> None:
    txt = paragraph.text
    if "«" not in txt:
        return

    def repl(m):
        key = m.group(1)
        return mapping.get(key, m.group(0))

    new_txt = PLACEHOLDER_RE.sub(repl, txt)
    if new_txt != txt:
        # Важно: paragraph.text = ... создаёт один run (упрощает форматирование),
        # но для вашего шаблона это ок, а шрифт мы всё равно фиксируем в конце.
        paragraph.text = new_txt


def replace_all(doc: Document, mapping: Dict[str, str]) -> None:
    for p in iter_paragraphs(doc):
        replace_placeholders_in_paragraph(p, mapping)


def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)
    paragraph._p = paragraph._element = None  # type: ignore


def delete_table_row(row: _Row) -> None:
    tr = row._tr
    tr.getparent().remove(tr)


def _sanitize_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r'[\\/:*?"<>|]+', "_", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


# ---------------- ПОИСК ГРАНИЦ БЛОКОВ ----------------

def _find_prosu_sud_index(doc: Document) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip().lower()
        if "прошу суд" in t:
            return i
    return None


def _find_common_block_index(doc: Document) -> Optional[int]:
    """
    Находим начало общего текста, который должен ОСТАТЬСЯ:
    "По всем вышеуказанным договорам ..."
    """
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip().lower()
        if "по всем вышеуказанным договорам" in t:
            return i
    return None


def _limit_for_defendant_sections(doc: Document) -> int:
    """
    Верхняя граница, до которой могут жить/удаляться договорные секции.
    Мы НЕ должны удалять общий блок ("По всем вышеуказанным договорам...") :contentReference[oaicite:2]{index=2}
    и тем более не должны залезать в "Прошу суд".
    """
    prosu = _find_prosu_sud_index(doc)
    common = _find_common_block_index(doc)

    limit = len(doc.paragraphs)
    if prosu is not None:
        limit = min(limit, prosu)
    if common is not None:
        limit = min(limit, common)

    return limit


def find_defendant_section_ranges(doc: Document) -> List[Tuple[int, int, int]]:
    """
    Находим диапазоны параграфов, соответствующие каждому ответчику (0..9),
    но только в "договорной" части, НЕ включая общий блок и НЕ включая "Прошу суд".
    """
    limit = _limit_for_defendant_sections(doc)

    starts: List[Tuple[int, int]] = []
    for i, p in enumerate(doc.paragraphs[:limit]):
        t = p.text or ""
        if "В соответствии с Договором" in t and "«Тип_договора" in t:
            m = re.search(r"«Тип_договора(\d*)»", t)
            if not m:
                continue
            suf = m.group(1)
            idx = 0 if suf == "" else int(suf)
            starts.append((idx, i))

    starts.sort(key=lambda x: x[1])

    ranges: List[Tuple[int, int, int]] = []
    for k, (idx, s) in enumerate(starts):
        e = starts[k + 1][1] if k + 1 < len(starts) else limit
        ranges.append((idx, s, e))

    ranges = [r for r in ranges if 0 <= r[0] <= 9]
    ranges.sort(key=lambda x: x[0])
    return ranges


# ---------------- ЛОГИКА "КТО ЕСТЬ" ----------------

def build_present_mask(mapping: Dict[str, str]) -> List[bool]:
    present: List[bool] = []
    for i in range(10):
        suf = "" if i == 0 else str(i)
        fio_val = (mapping.get("ФИО" + suf, "") or "").strip()
        present.append(bool(fio_val))
    if not any(present):
        present[0] = True
    return present


def _is_vzyskat_paragraph(p: Paragraph) -> bool:
    t = (p.text or "").strip()
    if not t:
        return False
    return re.match(r"^(?:\d+\s*[\.\)]\s*)?Взыскать с", t) is not None


def prune_unused_defendants(doc: Document, present: List[bool]) -> None:
    """
    Удаляет:
      - строки таблицы ОтветчикX для отсутствующих
      - секции договоров для отсутствующих (но НЕ трогает общий блок "По всем вышеуказанным договорам...") :contentReference[oaicite:3]{index=3}
      - пункты "Взыскать с ..." для отсутствующих
    """
    present = (present + [False] * 10)[:10]

    # 1) Таблица ответчиков (обычно первая)
    if doc.tables:
        table = doc.tables[0]
        defendant_rows = []
        for r_i, row in enumerate(table.rows):
            c0 = row.cells[0].text.strip()
            m = re.match(r"Ответчик(\d+):", c0)
            if m:
                num = int(m.group(1))  # 1..10
                defendant_rows.append((num, r_i))

        for num, r_i in sorted(defendant_rows, key=lambda x: x[1], reverse=True):
            idx = num - 1
            if 0 <= idx <= 9 and not present[idx]:
                delete_table_row(table.rows[r_i])

    # 2) Секции договоров — удаляем только внутри "договорной" зоны
    ranges = find_defendant_section_ranges(doc)
    to_delete = [r for r in ranges if 0 <= r[0] <= 9 and not present[r[0]]]

    # удаляем параграфы с конца
    paras = list(doc.paragraphs)
    for _, start, end in sorted(to_delete, key=lambda x: x[1], reverse=True):
        for p in paras[start:end][::-1]:
            delete_paragraph(p)

    # 3) Пункты "Взыскать с ..." — удаляем только отсутствующих
    req_paras = [p for p in doc.paragraphs if _is_vzyskat_paragraph(p)]
    for i in range(min(10, len(req_paras)) - 1, -1, -1):
        if not present[i]:
            delete_paragraph(req_paras[i])


# ---------------- ОСНОВНОЕ ----------------

@dataclass
class Inputs:
    excel_path: Path
    template_path: Path
    out_dir: Path  # выбранная пользователем "корневая" папка


def build_mapping(rows: pd.DataFrame) -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    for i, (_, r) in enumerate(rows.iterrows()):
        suf = "" if i == 0 else str(i)
        for col in rows.columns:
            key = normalize_key(col) + suf
            mapping[key] = safe_str(r[col])
    return mapping


def make_output_name(rows: pd.DataFrame, group_index: int) -> str:
    # имя по столбцу "Иск" у первого ответчика
    desired = safe_str(rows.iloc[0].get("Иск", ""))
    desired = _sanitize_filename(desired)

    if desired:
        return desired if desired.lower().endswith(".docx") else desired + ".docx"

    # fallback
    fio = safe_str(rows.iloc[0].get("ФИО", "")) or f"Группа_{group_index}"
    vnd = safe_str(rows.iloc[0].get("ВНД", ""))
    vnd = re.sub(r"[^\w\-]+", "_", vnd)[:20]
    fio_clean = _sanitize_filename(fio)[:60].strip("_")
    parts = [f"{group_index:03d}", fio_clean]
    if vnd:
        parts.append(vnd)
    return "Иск_" + "_".join([p for p in parts if p]) + ".docx"


def generate_claims(inputs: Inputs) -> int:
    df = pd.read_excel(inputs.excel_path)
    if df.empty:
        raise ValueError("Excel пустой — нечего формировать.")

    # --- создаём подпапку по названию Excel-файла ---
    excel_folder_name = _sanitize_filename(inputs.excel_path.stem) or "Иски"
    out_dir = inputs.out_dir / excel_folder_name
    out_dir.mkdir(parents=True, exist_ok=True)

    # --- Реестр: новые столбцы ---
    if "ВНД с чем объединены" not in df.columns:
        df["ВНД с чем объединены"] = ""
    if "Файл_иска" not in df.columns:
        df["Файл_иска"] = ""

    total = 0
    group_index = 1

    for start in range(0, len(df), 10):
        chunk = df.iloc[start:start + 10].copy()
        chunk_index = chunk.index

        doc = Document(inputs.template_path)
        mapping = build_mapping(chunk)

        # 1) Подстановка
        replace_all(doc, mapping)

        # 2) Удаляем пустых ответчиков/секции/пункты "Взыскать с"
        present = build_present_mask(mapping)
        if not all(present):
            prune_unused_defendants(doc, present)

        # 3) Имя файла иска — по столбцу "Иск" 1-го ответчика
        out_name = make_output_name(chunk, group_index)
        out_path = out_dir / out_name

        # 4) Шрифт
        enforce_times_new_roman_12(doc)

        # 5) Сохранение
        doc.save(out_path)

        # 6) Реестр
        first_vnd = safe_str(chunk.iloc[0].get("ВНД", ""))
        df.loc[chunk_index, "ВНД с чем объединены"] = first_vnd
        df.loc[chunk_index, "Файл_иска"] = out_name

        total += 1
        group_index += 1

    # сохраняем реестр внутрь созданной подпапки
    registry_path = out_dir / "Реестр_исков.xlsx"
    df.to_excel(registry_path, index=False)

    return total


# ---------------- GUI ----------------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Иски-бот (Excel → Word)")
        self.resizable(False, False)

        self.excel_path: Optional[Path] = None
        self.template_path: Optional[Path] = None
        self.out_dir: Optional[Path] = None

        pad = {"padx": 10, "pady": 6}

        self.lbl1 = tk.Label(self, text="1) Excel с данными:")
        self.lbl1.grid(row=0, column=0, sticky="w", **pad)
        self.btn_excel = tk.Button(self, text="Загрузить Excel…", command=self.pick_excel, width=20)
        self.btn_excel.grid(row=0, column=1, **pad)
        self.val_excel = tk.Label(self, text="—", fg="#555", wraplength=420, justify="left")
        self.val_excel.grid(row=1, column=0, columnspan=2, sticky="w", **pad)

        self.lbl2 = tk.Label(self, text="2) Word-шаблон:")
        self.lbl2.grid(row=2, column=0, sticky="w", **pad)
        self.btn_docx = tk.Button(self, text="Загрузить Word…", command=self.pick_docx, width=20)
        self.btn_docx.grid(row=2, column=1, **pad)
        self.val_docx = tk.Label(self, text="—", fg="#555", wraplength=420, justify="left")
        self.val_docx.grid(row=3, column=0, columnspan=2, sticky="w", **pad)

        self.lbl3 = tk.Label(self, text="3) Корневая папка (в ней создастся подпапка по имени Excel):")
        self.lbl3.grid(row=4, column=0, sticky="w", **pad)
        self.btn_out = tk.Button(self, text="Выбрать папку…", command=self.pick_outdir, width=20)
        self.btn_out.grid(row=4, column=1, **pad)
        self.val_out = tk.Label(self, text="—", fg="#555", wraplength=420, justify="left")
        self.val_out.grid(row=5, column=0, columnspan=2, sticky="w", **pad)

        self.btn_run = tk.Button(self, text="Сформировать иски", command=self.run, width=30)
        self.btn_run.grid(row=6, column=0, columnspan=2, pady=12)

        self.attributes("-topmost", True)
        self.after(350, lambda: self.attributes("-topmost", False))

    def pick_excel(self):
        path = filedialog.askopenfilename(
            title="Выберите Excel",
            filetypes=[("Excel files", "*.xlsx *.xls *.xlsm *.xlsb"), ("All files", "*.*")]
        )
        if path:
            self.excel_path = Path(path)
            self.val_excel.config(text=str(self.excel_path))

    def pick_docx(self):
        path = filedialog.askopenfilename(
            title="Выберите Word-шаблон",
            filetypes=[("Word files", "*.docx *.docm"), ("All files", "*.*")]
        )
        if path:
            self.template_path = Path(path)
            self.val_docx.config(text=str(self.template_path))

    def pick_outdir(self):
        path = filedialog.askdirectory(title="Выберите корневую папку")
        if path:
            self.out_dir = Path(path)
            self.val_out.config(text=str(self.out_dir))

    def run(self):
        if not self.excel_path or not self.template_path or not self.out_dir:
            messagebox.showwarning("Не хватает данных", "Выберите Excel, Word-шаблон и папку сохранения.")
            return
        try:
            total = generate_claims(Inputs(self.excel_path, self.template_path, self.out_dir))
            subfolder = _sanitize_filename(self.excel_path.stem) or "Иски"
            out_final = self.out_dir / subfolder
            messagebox.showinfo(
                "Готово",
                f"Сформировано файлов: {total}\nПапка: {out_final}\nСоздан реестр: Реестр_исков.xlsx"
            )
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))


def main():
    if tk is None:
        raise RuntimeError("Tkinter не доступен в этой среде. Запустите на Windows/macOS/Linux с GUI.")
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()
