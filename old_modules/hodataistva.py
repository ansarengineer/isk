# -*- coding: utf-8 -*-
"""
Ходатайства-бот: формирует ходатайства из Excel по Word-шаблону.

Требования:
- Папка результата ВСЕГДА называется "Ходатайство"
- Название файла берётся из столбца 128 "Ходатайство"
- «Дата_составления_иска» = дата создания документа
- Весь текст = Times New Roman, 12 pt
"""

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Tuple

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.oxml.ns import qn

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except Exception:
    tk = None


PLACEHOLDER_RE = re.compile(r"«([^»]+)»")


# ---------- utils ----------

def normalize_key(col_name: str) -> str:
    return re.sub(r"\s+", "_", str(col_name).strip())


def safe_str(v) -> str:
    if pd.isna(v):
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v).strip()


def sanitize_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]+', "_", (s or "").strip())
    return re.sub(r"\s+", " ", s).strip()


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p


def replace_placeholders(doc: Document, mapping: Dict[str, str]) -> None:
    for p in iter_paragraphs(doc):
        txt = p.text
        if "«" not in txt:
            continue

        def repl(m):
            return mapping.get(m.group(1), m.group(0))

        new_txt = PLACEHOLDER_RE.sub(repl, txt)
        if new_txt != txt:
            p.text = new_txt


def enforce_tnr_12(doc: Document) -> None:
    def apply(p: Paragraph):
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            for k in ("ascii", "hAnsi", "cs", "eastAsia"):
                rFonts.set(qn(f"w:{k}"), "Times New Roman")

    for p in doc.paragraphs:
        apply(p)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    apply(p)


def build_mapping(row: pd.Series) -> Dict[str, str]:
    return {normalize_key(c): safe_str(row[c]) for c in row.index}


def make_filename(row: pd.Series, idx: int) -> str:
    name = safe_str(row.get("Ходатайство", ""))  # столбец 128
    name = sanitize_filename(name)
    if name:
        return name if name.lower().endswith(".docx") else name + ".docx"

    fio = sanitize_filename(safe_str(row.get("ФИО", "")) or f"Ответчик_{idx+1}")
    return f"Ходатайство_{fio}.docx"


# ---------- core ----------

@dataclass
class Inputs:
    excel: Path
    template: Path
    out_root: Path


def generate(inputs: Inputs) -> Tuple[int, Path]:
    df = pd.read_excel(inputs.excel)
    if df.empty:
        raise ValueError("Excel пустой")

    out_dir = inputs.out_root / "Ходатайство"
    out_dir.mkdir(parents=True, exist_ok=True)

    today = datetime.now().strftime("%d.%m.%Y")
    count = 0

    for i, (_, row) in enumerate(df.iterrows()):
        doc = Document(inputs.template)
        mapping = build_mapping(row)
        mapping["Дата_составления_иска"] = today

        replace_placeholders(doc, mapping)
        enforce_tnr_12(doc)

        filename = make_filename(row, i)
        doc.save(out_dir / filename)
        count += 1

    return count, out_dir


# ---------- GUI ----------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Ходатайства-бот")
        self.resizable(False, False)

        self.excel: Optional[Path] = None
        self.template: Optional[Path] = None
        self.out_root: Optional[Path] = None

        pad = {"padx": 10, "pady": 6}

        tk.Label(self, text="1) Excel файл:").grid(row=0, column=0, sticky="w", **pad)
        tk.Button(self, text="Загрузить Excel…", command=self.pick_excel, width=25).grid(row=0, column=1, **pad)
        self.l_excel = tk.Label(self, text="—", fg="#555", wraplength=500)
        self.l_excel.grid(row=1, column=0, columnspan=2, sticky="w", **pad)

        tk.Label(self, text="2) Шаблон ходатайства:").grid(row=2, column=0, sticky="w", **pad)
        tk.Button(self, text="Загрузить шаблон…", command=self.pick_template, width=25).grid(row=2, column=1, **pad)
        self.l_tpl = tk.Label(self, text="—", fg="#555", wraplength=500)
        self.l_tpl.grid(row=3, column=0, columnspan=2, sticky="w", **pad)

        tk.Label(self, text="3) Куда сохранить:").grid(row=4, column=0, sticky="w", **pad)
        tk.Button(self, text="Выбрать папку…", command=self.pick_out, width=25).grid(row=4, column=1, **pad)
        self.l_out = tk.Label(self, text="—", fg="#555", wraplength=500)
        self.l_out.grid(row=5, column=0, columnspan=2, sticky="w", **pad)

        tk.Button(self, text="Сформировать ходатайства", command=self.run, width=36)\
            .grid(row=6, column=0, columnspan=2, pady=12)

        self.attributes("-topmost", True)
        self.after(300, lambda: self.attributes("-topmost", False))

    def pick_excel(self):
        p = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx *.xls")])
        if p:
            self.excel = Path(p)
            self.l_excel.config(text=p)

    def pick_template(self):
        p = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if p:
            self.template = Path(p)
            self.l_tpl.config(text=p)

    def pick_out(self):
        p = filedialog.askdirectory()
        if p:
            self.out_root = Path(p)
            self.l_out.config(text=p)

    def run(self):
        if not all([self.excel, self.template, self.out_root]):
            messagebox.showwarning("Ошибка", "Выберите Excel, шаблон и папку")
            return
        try:
            total, folder = generate(Inputs(self.excel, self.template, self.out_root))
            messagebox.showinfo("Готово", f"Создано ходатайств: {total}\nПапка: {folder}")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))


def main():
    if tk is None:
        raise RuntimeError("Tkinter недоступен")
    App().mainloop()


if __name__ == "__main__":
    main()
