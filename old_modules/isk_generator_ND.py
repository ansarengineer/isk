import os
import re
import pandas as pd
from docx import Document

import tkinter as tk
from tkinter import ttk, filedialog, messagebox


# --- Явные соответствия: плейсхолдер (внутри « ») -> колонка Excel ---
ALIAS_MAP = {
    # Ваш кейс:
    "БП_Дата_уступки_г.": "БП/3П дата уступки г.",
}

# ---------- DOCX helpers ----------
PLACEHOLDER_RE = re.compile(r"«([^»]+)»")


def normalize_key(s: str) -> str:
    s = str(s).strip().lower()
    s = s.replace("ё", "е")
    s = re.sub(r"[_\s]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_placeholders_in_paragraph(paragraph, mapping: dict):
    # Надёжная замена даже если «ПЛЕЙСХОЛДЕР» разбит на несколько runs
    if not paragraph.runs:
        return

    full_text = "".join(r.text for r in paragraph.runs)
    if "«" not in full_text:
        return

    def repl(m):
        key = m.group(1)
        return str(mapping.get(key, m.group(0)))  # если нет значения — оставляем плейсхолдер

    new_text = PLACEHOLDER_RE.sub(repl, full_text)
    if new_text == full_text:
        return

    first = paragraph.runs[0]
    for r in paragraph.runs:
        r.text = ""
    first.text = new_text


def extract_placeholders(doc: Document) -> set:
    found = set()
    for p in iter_all_paragraphs(doc):
        text = "".join(r.text for r in p.runs) if p.runs else p.text
        for m in PLACEHOLDER_RE.finditer(text):
            found.add(m.group(1))
    return found


def safe_filename(s: str) -> str:
    s = str(s).strip()
    s = re.sub(r"[\\/:*?\"<>|]+", "_", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:180] if len(s) > 180 else s


def build_placeholder_mapping(row: pd.Series, placeholders: set, col_norm_map: dict) -> dict:
    mapping = {}
    for ph in placeholders:
        excel_col = ALIAS_MAP.get(ph)
        if excel_col is None:
            ph_norm = normalize_key(ph)
            excel_col = col_norm_map.get(ph_norm)

        if not excel_col:
            continue

        val = row.get(excel_col, "")
        if pd.isna(val):
            val = ""
        mapping[ph] = val
    return mapping


def ensure_folder(base_dir: str, folder_name: str) -> str:
    """
    Создаёт папку base_dir/folder_name.
    Если уже есть — создаёт base_dir/folder_name (2), (3), ...
    """
    target = os.path.join(base_dir, folder_name)
    if not os.path.exists(target):
        os.makedirs(target, exist_ok=True)
        return target

    i = 2
    while True:
        candidate = f"{target} ({i})"
        if not os.path.exists(candidate):
            os.makedirs(candidate, exist_ok=True)
            return candidate
        i += 1


def detect_isk_column(df: pd.DataFrame):
    """
    Требование: имя иска брать из 122 столбца "Иск".
    Поддерживаем:
      - колонку с названием "Иск"
      - если нет названия, берём 122-й столбец по индексу (121)
    """
    # 1) по названию
    for c in df.columns:
        if normalize_key(c) == normalize_key("Иск"):
            return c

    # 2) по позиции (122-й столбец)
    if len(df.columns) >= 122:
        return df.columns[121]

    return None


# ---------- Main generation ----------
def generate_docs(excel_path: str, template_path: str, base_out_dir: str, log_fn):
    if not os.path.isfile(excel_path):
        raise FileNotFoundError(f"Excel не найден: {excel_path}")
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Шаблон не найден: {template_path}")
    if not base_out_dir:
        raise ValueError("Не указана папка для сохранения.")

    # 1) Создаём папку "Иски НД" внутри указанного пути
    claims_dir = ensure_folder(base_out_dir, "Иски НД")
    log_fn(f"📁 Папка для исков: {claims_dir}")

    # 2) Читаем Excel
    df = pd.read_excel(excel_path)
    if df.empty:
        raise ValueError("Excel пустой.")

    # 3) Определяем колонку для имени файла (Иск)
    isk_col = detect_isk_column(df)
    if not isk_col:
        raise ValueError('Не найдена колонка "Иск" и нет 122-го столбца (позиция 121).')

    # Нормализованная карта колонок Excel (для авто-сопоставления плейсхолдеров)
    col_norm_map = {normalize_key(c): c for c in df.columns}

    # 4) Плейсхолдеры из шаблона
    template_doc = Document(template_path)
    placeholders = extract_placeholders(template_doc)
    log_fn(f"Найдено плейсхолдеров в шаблоне: {len(placeholders)}")

    created = 0
    for idx, row in df.iterrows():
        doc = Document(template_path)

        mapping = build_placeholder_mapping(row, placeholders, col_norm_map)
        for p in iter_all_paragraphs(doc):
            replace_placeholders_in_paragraph(p, mapping)

        # 5) Имя файла — строго из колонки "Иск" (122-й столбец)
        raw_name = row.get(isk_col, "")
        if pd.isna(raw_name) or str(raw_name).strip() == "":
            raw_name = f"иск_{idx+1}"

        file_base = safe_filename(raw_name)
        out_path = os.path.join(claims_dir, f"{file_base}.docx")

        # защита от перезаписи: добавим (2), (3), ...
        if os.path.exists(out_path):
            n = 2
            while True:
                alt = os.path.join(claims_dir, f"{file_base} ({n}).docx")
                if not os.path.exists(alt):
                    out_path = alt
                    break
                n += 1

        doc.save(out_path)
        created += 1
        log_fn(f"✔ Создано: {os.path.basename(out_path)}")

    return created, claims_dir


# ---------- GUI ----------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Генератор исков из Excel")
        self.geometry("780x440")
        self.resizable(True, True)

        self.excel_path = tk.StringVar()
        self.template_path = tk.StringVar()
        self.out_dir = tk.StringVar()

        self._build_ui()

    def _build_ui(self):
        frm = ttk.Frame(self, padding=12)
        frm.pack(fill="both", expand=True)

        ttk.Label(frm, text="1) Excel с данными:").grid(row=0, column=0, sticky="w")
        ttk.Entry(frm, textvariable=self.excel_path).grid(row=1, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(frm, text="Выбрать Excel", command=self.pick_excel).grid(row=1, column=1, sticky="ew")

        ttk.Label(frm, text="2) Шаблон иска (.docx):").grid(row=2, column=0, sticky="w", pady=(10, 0))
        ttk.Entry(frm, textvariable=self.template_path).grid(row=3, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(frm, text="Выбрать шаблон", command=self.pick_template).grid(row=3, column=1, sticky="ew")

        ttk.Label(frm, text="3) Путь (базовая папка) для сохранения:").grid(row=4, column=0, sticky="w", pady=(10, 0))
        ttk.Entry(frm, textvariable=self.out_dir).grid(row=5, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(frm, text="Выбрать папку", command=self.pick_outdir).grid(row=5, column=1, sticky="ew")

        ttk.Button(frm, text="Сформировать иски", command=self.run, style="Accent.TButton").grid(
            row=6, column=0, columnspan=2, sticky="ew", pady=(12, 8)
        )

        ttk.Label(frm, text="Лог:").grid(row=7, column=0, sticky="w")
        self.log = tk.Text(frm, height=12, wrap="word")
        self.log.grid(row=8, column=0, columnspan=2, sticky="nsew")

        frm.columnconfigure(0, weight=1)
        frm.rowconfigure(8, weight=1)

        style = ttk.Style(self)
        try:
            style.configure("Accent.TButton", font=("Segoe UI", 10, "bold"))
        except Exception:
            pass

    def log_write(self, msg: str):
        self.log.insert("end", msg + "\n")
        self.log.see("end")
        self.update_idletasks()

    def pick_excel(self):
        p = filedialog.askopenfilename(
            title="Выберите Excel файл",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if p:
            self.excel_path.set(p)

    def pick_template(self):
        p = filedialog.askopenfilename(
            title="Выберите шаблон .docx",
            filetypes=[("Word files", "*.docx")]
        )
        if p:
            self.template_path.set(p)

    def pick_outdir(self):
        p = filedialog.askdirectory(title="Выберите базовую папку для сохранения")
        if p:
            self.out_dir.set(p)

    def run(self):
        excel_path = self.excel_path.get().strip()
        template_path = self.template_path.get().strip()
        base_out_dir = self.out_dir.get().strip()

        if not excel_path or not template_path or not base_out_dir:
            messagebox.showwarning("Не хватает данных", "Выберите Excel, шаблон и папку сохранения.")
            return

        self.log_write("=== Старт ===")
        try:
            count, claims_dir = generate_docs(excel_path, template_path, base_out_dir, self.log_write)
            self.log_write(f"=== Готово. Создано документов: {count} ===")
            self.log_write(f"Папка с исками: {claims_dir}")
            messagebox.showinfo("Готово", f"Создано документов: {count}\nПапка: {claims_dir}")
        except Exception as e:
            self.log_write(f"✖ Ошибка: {e}")
            messagebox.showerror("Ошибка", str(e))


if __name__ == "__main__":
    App().mainloop()
